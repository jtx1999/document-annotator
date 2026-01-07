import streamlit as st
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from typing import List, Dict, Tuple
from google import genai
from pydantic import BaseModel, Field
from io import BytesIO


class AnswerKey(BaseModel):
    para_id: int = Field(..., description="Paragraph ID where the answer key is located")
    answer: str = Field(..., description="The extracted answer key text")


class AnswerKeys(BaseModel):
    answer_keys: List[AnswerKey] = Field(..., description="List of identified answer keys")


def table_to_markdown(table: Table) -> str:
    """Converts a docx Table object to a Markdown string."""
    md_rows = []
    for row in table.rows:
        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        md_rows.append(f"| {' | '.join(cells)} |")
    
    # Create the header separator
    if len(md_rows) > 0:
        header_sep = f"| {' | '.join(['---'] * len(table.columns))} |"
        md_rows.insert(1, header_sep)
    
    return "\n".join(md_rows)


def get_document_content(file_path: str) -> Tuple[Document, List[Dict[str, str]]]:
    """
    Extracts text from a .docx file and returns a list of indexed paragraphs.
    """
    doc = Document(file_path)
    content = []

    i = 0
    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            if para.text.strip():
                content.append({
                    "type": "paragraph",
                    "content": para.text.strip(),
                    "para_id": i,
                })
            i += 1
        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            table_md = table_to_markdown(table)
            content.append({
                "type": "table",
                "content": table_md
            })
    return doc, content

prompt = """
You are an expert at identifying answer keys in educational documents.
Given the following document content, identify the paragraphs containing questions, and their corresponding answer keys.
Output a list of objects with 'para_id' and 'answer' fields. The `para_id` corresponds to the index of the QUESTION in the exam paper, and the `answer` is the text of the answer key.
Some question may span multiple paragraphs, but the `para_id` for that answer should point to the beginning of the question.
Some answers may correpond to multiple questions, especially for multiple choice questions. In such cases, list each question's `para_id` separately with the corresponding answer to that question.
For questions that contain sub-questions, break down the answer for each sub-question and list the `para_id` of each sub-question.
There might be multiple exams in the document. Identify the answer keys for all exams present.
"""


def identify_answer_keys(doc_content: List[Dict[str, str]]) -> AnswerKeys:
    client = genai.Client(api_key=st.secrets["GENAI_API_KEY"])
    response = client.models.generate_content(
        model="gemini-3-flash-preview",
        contents=prompt + "\nDocument Content:\n" + str(doc_content),
        config={
            "response_mime_type": "application/json",
            "response_json_schema": AnswerKeys.model_json_schema(),
        }
    )
    return AnswerKeys.model_validate_json(response.text)


def add_comments(doc: Document, answer_keys: AnswerKeys):
    for answer_key in answer_keys.answer_keys:
        paragraph = doc.paragraphs[answer_key.para_id]
        doc.add_comment(paragraph.runs, answer_key.answer, author="ChemistryAI")
    return doc


def main():
    st.set_page_config(
        page_title="文档答案标注工具",
        page_icon="📄",
        layout="centered",
    )
    st.title("文档答案标注工具")
    st.write("上传一个 .docx 文件以对其进行答案标注。")

    uploaded_file = st.file_uploader("上传一个 .docx 文件", type=["docx"])
    if uploaded_file:
        file_name = uploaded_file.name
        doc = Document(uploaded_file)
        _, doc_content = get_document_content(uploaded_file)

        with st.spinner("正在识别答案..."):
            answer_keys = identify_answer_keys(doc_content)

        with st.spinner("正在标注文档..."):
            annotated_doc = add_comments(doc, answer_keys)

        st.success("文档标注成功！")

        # Save the annotated document to a BytesIO object for download
        annotated_file = BytesIO()
        annotated_doc.save(annotated_file)
        annotated_file.seek(0)

        st.download_button(
            label="下载标注后的文档",
            data=annotated_file,
            file_name="annotated_" + file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            on_click="ignore",
        )


if __name__ == "__main__":
    main()
